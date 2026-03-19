# -*- coding: utf-8 -*-
"""
會議紀錄助理 — Streamlit 版（企業級 SaaS UI）

【功能說明】
1. 會議記錄內容支援列號顯示（Ace Editor）
2. 工具列固定在內容下方（排版 / 搜尋 / 附件 / 預覽 / Word）
3. 預覽使用浮動視窗（dialog）
4. 支援附件列數、搜尋定位、Word 匯出
5. 三欄版面（基本資料 / 內容 / 執行者）

⚠️ 注意：
這段為說明文字，已用 docstring 包起來，避免 Python 語法錯誤
"""

import os
import io
import re
import math
import base64
import tempfile

import streamlit as st

# ===== docx =====
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ===== Ace Editor（列號顯示）=====

try:
    from streamlit_ace import st_ace
    HAS_ACE = True
except ImportError:
    HAS_ACE = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False


TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")
P1_ROWS = 11
P2_ROWS = 26
MAX_CHARS = 30
DELIMITER_DISPLAY = "/"
DELIMITER_ALIASES = ["/", "#", "|"]

# ── 附件佔位符格式 ────────────────────────────────────────────────
# 每一個附件在 editor 中佔 row_cost 行，每行都是相同的佔位符文字
# 格式：  <<<ATT:標籤文字>>>
# 用三個角括號確保不會和正常內容衝突
_ATT_PREFIX = "<<<ATT:"
_ATT_SUFFIX = ">>>"


def make_placeholder(label: str) -> str:
    """產生單行附件佔位符文字"""
    return f"{_ATT_PREFIX}{label}{_ATT_SUFFIX}"


def is_placeholder(line: str) -> bool:
    s = line.strip()
    return s.startswith(_ATT_PREFIX) and s.endswith(_ATT_SUFFIX)


def placeholder_label(line: str) -> str:
    s = line.strip()
    return s[len(_ATT_PREFIX):-len(_ATT_SUFFIX)]



def build_editor_value(content_text: str, attachments: list) -> str:
    """向後相容：導向新的 build_editor_display"""
    return build_editor_display(content_text, attachments)


def strip_placeholders(editor_text: str, expected_placeholders: list = None) -> tuple:
    """
    從 editor 回傳的文字還原純文字：
    - 移除附件佔位符行（<<<ATT:...>>>）
    - 核對佔位符數量，偵測竄改
    回傳 (純文字, 是否竄改)
    """
    lines = editor_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    clean = []
    tampered = False
    found_labels: dict = {}

    for line in lines:
        stripped = line.strip()
        if is_placeholder(stripped):
            lbl = placeholder_label(stripped)
            found_labels[lbl] = found_labels.get(lbl, 0) + 1
            continue
        # 部分破壞的佔位符標記 → 竄改
        if _ATT_PREFIX in line or _ATT_SUFFIX in line:
            tampered = True
            continue
        clean.append(line)

    if expected_placeholders is not None:
        expected: dict = {}
        for label, count in expected_placeholders:
            expected[label] = expected.get(label, 0) + count
        if found_labels != expected:
            tampered = True

    return "\n".join(clean), tampered


def get_expected_placeholders(attachments: list) -> list:
    """回傳 [(label, row_cost), ...] 供核對用"""
    return [(att.get("label", "附件"), int(att.get("row_cost", 1))) for att in attachments]


def split_content_owner_line(line: str):
    for sep in DELIMITER_ALIASES:
        if sep in line:
            left, right = line.split(sep, 1)
            return left, right, sep
    return None


MAX_UNITS = 60  # 每列上限：60半形單位 = 30中文字

COLOR_NAME_TO_HEX = {
    "red": "C00000",
    "blue": "1F4E79",
    "green": "548235",
    "orange": "C55A11",
    "purple": "7030A0",
    "gray": "666666",
    "black": "000000",
}
COLOR_TAG_ALIASES = {
    "紅": "red", "红": "red", "red": "red",
    "藍": "blue", "蓝": "blue", "blue": "blue",
    "綠": "green", "绿": "green", "green": "green",
    "橘": "orange", "橙": "orange", "orange": "orange",
    "紫": "purple", "purple": "purple",
    "灰": "gray", "gray": "gray", "grey": "gray",
    "黑": "black", "black": "black",
}
COLOR_CANONICAL_TAG = {
    "red": "紅",
    "blue": "藍",
    "green": "綠",
    "orange": "橘",
    "purple": "紫",
    "gray": "灰",
    "black": "黑",
}


def _char_units(ch: str) -> int:
    """CJK全形=2單位，英數半形=1單位"""
    cp = ord(ch)
    if (0x4E00 <= cp <= 0x9FFF or 0x3400 <= cp <= 0x4DBF or
            0xF900 <= cp <= 0xFAFF or 0x3000 <= cp <= 0x303F or
            0x3040 <= cp <= 0x30FF or 0xAC00 <= cp <= 0xD7AF or
            0x2E80 <= cp <= 0x2EFF or
            0xFF01 <= cp <= 0xFF60 or 0xFFE0 <= cp <= 0xFFE6):
        return 2
    return 1


def split_styled_segments(text: str):
    import re
    src = "" if text is None else str(text)
    pattern = re.compile(r"\[(\/)?([^\]]*)\]")
    pos = 0
    stack = [None]
    out = []

    def push_text(chunk):
        if not chunk:
            return
        color = stack[-1]
        if out and out[-1][1] == color:
            out[-1] = (out[-1][0] + chunk, color)
        else:
            out.append((chunk, color))

    for m in pattern.finditer(src):
        push_text(src[pos:m.start()])
        is_close = bool(m.group(1))
        raw_name = (m.group(2) or "").strip().lower()
        if is_close:
            if len(stack) > 1:
                if raw_name:
                    target = COLOR_TAG_ALIASES.get(raw_name)
                    while len(stack) > 1 and stack[-1] != target:
                        stack.pop()
                    if len(stack) > 1:
                        stack.pop()
                else:
                    stack.pop()
            else:
                push_text(m.group(0))
        else:
            key = COLOR_TAG_ALIASES.get(raw_name)
            if key:
                stack.append(key)
            else:
                push_text(m.group(0))
        pos = m.end()
    push_text(src[pos:])
    return out


def segments_to_plain(segments) -> str:
    return "".join(txt for txt, _ in segments)


def segments_to_tagged(segments) -> str:
    parts = []
    current = None
    for txt, color in segments:
        if color != current:
            if current is not None:
                parts.append("[/]")
            if color is not None:
                parts.append(f"[{COLOR_CANONICAL_TAG.get(color, color)}]")
            current = color
        parts.append(txt)
    if current is not None:
        parts.append("[/]")
    return "".join(parts)


def segments_to_html(segments) -> str:
    html_parts = []
    for txt, color in segments:
        safe = (txt or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace(" ", "&nbsp;").replace("\n", "<br>")
        if color:
            html_parts.append(f"<span style='color:#{COLOR_NAME_TO_HEX.get(color, '000000')}'>{safe}</span>")
        else:
            html_parts.append(safe)
    return "".join(html_parts)


def styled_chars_from_text(text: str):
    chars = []
    for chunk, color in split_styled_segments(text):
        for ch in chunk:
            chars.append((ch, color))
    return chars


def wrap_styled_line_30(line: str, initial_units: int = 0, subsequent_units: int = 0):
    chars = styled_chars_from_text(("" if line is None else str(line)).rstrip("\r"))
    if not chars:
        return [[]]

    tokens = []
    i = 0
    n = len(chars)
    while i < n:
        ch, color = chars[i]
        if ch.isascii() and ch.isalnum():
            token = []
            while i < n and chars[i][0].isascii() and chars[i][0].isalnum():
                token.append(chars[i])
                i += 1
            tokens.append(token)
        else:
            tokens.append([chars[i]])
            i += 1

    limit = MAX_UNITS
    result = []
    cur = []
    cur_u = initial_units
    next_line_units = subsequent_units

    def token_units(tok):
        return sum(_char_units(ch) for ch, _ in tok)

    def append_line(tok_list):
        segs = []
        for ch, color in tok_list:
            if segs and segs[-1][1] == color:
                segs[-1] = (segs[-1][0] + ch, color)
            else:
                segs.append((ch, color))
        result.append(segs)

    for tok in tokens:
        tu = token_units(tok)
        if cur_u + tu > limit:
            if cur:
                append_line(cur)
                cur, cur_u = [], next_line_units
            if cur_u + tu > limit and tu > (limit - cur_u):
                buf = []
                buf_u = cur_u
                for item in tok:
                    cu = _char_units(item[0])
                    if buf_u + cu > limit and buf:
                        append_line(buf)
                        buf, buf_u = [item], next_line_units + cu
                    else:
                        buf.append(item)
                        buf_u += cu
                cur, cur_u = buf, buf_u
            else:
                cur, cur_u = tok[:], cur_u + tu
        else:
            cur.extend(tok)
            cur_u += tu

    if cur:
        append_line(cur)
    return result if result else [[]]


def wrap_plain_line_30(line: str, initial_units: int = 0, subsequent_units: int = 0) -> list:
    return [segments_to_plain(seg_line) for seg_line in wrap_styled_line_30(line, initial_units=initial_units, subsequent_units=subsequent_units)]


# ── 內容的設計原則 ────────────────────────────────────────────────
# session_state["content"] 存「目前編輯器中的純文字內容」（不含附件佔位符）。
# 可為：
# - 使用者原始輸入內容（尚未按 Word 排版）
# - 已按過 Word 排版後的顯示內容
#
# 因此：
# - 編輯器顯示時，不做即時自動換行/縮排
# - 需要正式套用 Word 排版時，統一先 normalize，再 expand
# - Word 輸出 / 列數統計 / 搜尋 都以 normalize 後的結果為準
# ─────────────────────────────────────────────────────────────────

INDENT = "　　"          # 公文首行縮排（全形空格×2 = 4 單位）
INDENT_UNITS = 4         # INDENT 佔用的單位數
CONT_MARK = "\u200b"   # 零寬字元：標記「自動換行產生的接續行」


def units_to_indent_text(units: int) -> str:
    """依單位數轉成可顯示縮排；2單位=1個全形空格，1單位=1個半形空格。"""
    units = max(0, int(units or 0))
    return "　" * (units // 2) + (" " if units % 2 else "")


def detect_outline_prefix(text: str):
    """
    回傳 (首行縮排單位, 編號字串)：
    - 一、      -> 前面縮 2 個全形空格（4 單位）
    - 1.       -> 前面縮 3 個全形空格（6 單位）
    - (1)/(（1）)-> 前面縮 4 個全形空格（8 單位）
    其餘一般段落沿用原本 2 個全形空格（4 單位），且無編號字串。
    """
    import re
    s = "" if text is None else str(text).lstrip("　 ")
    m = re.match(r'^([一二三四五六七八九十]+、)', s)
    if m:
        return 4, m.group(1)
    m = re.match(r'^(\d+\.)', s)
    if m:
        return 6, m.group(1)
    m = re.match(r'^((?:\(\d+\)|（\d+）))', s)
    if m:
        return 8, m.group(1)
    return INDENT_UNITS, ""


def detect_outline_indent_units(text: str) -> int:
    return detect_outline_prefix(text)[0]


def get_display_indent(text: str) -> str:
    return units_to_indent_text(detect_outline_indent_units(text))


def get_hanging_indent_units(text: str) -> int:
    """接續行對齊『內容起點』，不可超出編號。"""
    indent_units, marker = detect_outline_prefix(text)
    if marker:
        marker_units = sum(_char_units(ch) for ch in marker)
        return indent_units + marker_units
    return 0


def get_hanging_indent_text(text: str) -> str:
    return units_to_indent_text(get_hanging_indent_units(text))


def strip_display_indent(line: str) -> str:
    """移除 Word排版時首行加上的顯示縮排，還原成原始段落內容。"""
    s = "" if line is None else str(line)
    for fullwidth_spaces in (4, 3, 2):
        prefix = "　" * fullwidth_spaces
        rest = s[len(prefix):] if s.startswith(prefix) else None
        if rest is not None:
            units = detect_outline_indent_units(rest)
            if units == fullwidth_spaces * 2:
                return rest
    if s.startswith(INDENT):
        return s[len(INDENT):]
    return s


def strip_continuation_display_indent(prev_raw: str, line: str) -> str:
    """移除接續行為了避開編號而加上的可見縮排。"""
    s = "" if line is None else str(line)
    prefix = get_hanging_indent_text(prev_raw)
    if prefix and s.startswith(prefix):
        return s[len(prefix):]
    return s


def expand_content_lines(content_text: str) -> list:
    """
    將 content（原始輸入行）展開為「顯示用行列表」。
    每個元素是一個 dict，包含純文字顯示與帶色彩資訊的 segments。
    """
    src = [] if not content_text else str(content_text).replace("\r\n", "\n").replace("\r", "\n").split("\n")
    result = []
    owner_idx = 0

    for raw in src:
        if not raw.strip():
            result.append({
                "display": "",
                "display_tagged": "",
                "segments": [],
                "is_first": True,
                "owner_idx": owner_idx,
            })
            owner_idx += 1
            continue

        parsed = split_content_owner_line(raw)
        if parsed:
            left, right, sep = parsed
            left_segments = split_styled_segments(left)
            owner_segments = split_styled_segments(right.strip())
            result.append({
                "display": segments_to_plain(left_segments),
                "display_tagged": segments_to_tagged(left_segments),
                "segments": left_segments,
                "is_first": True,
                "owner_idx": owner_idx,
                "owner_inline": segments_to_plain(owner_segments),
                "owner_inline_segments": owner_segments,
                "is_resolution": True,
                "separator": sep,
            })
            owner_idx += 1
        else:
            first_indent = get_display_indent(raw)
            first_indent_units = detect_outline_indent_units(raw)
            hanging_indent = get_hanging_indent_text(raw)
            hanging_units = get_hanging_indent_units(raw)
            seg_lines = wrap_styled_line_30(raw, initial_units=first_indent_units, subsequent_units=hanging_units)
            for i, seg_line in enumerate(seg_lines):
                if i == 0:
                    display_segments = ([(first_indent, None)] + seg_line)
                else:
                    display_segments = ([(hanging_indent, None)] + seg_line) if hanging_indent else seg_line
                result.append({
                    "display": segments_to_plain(display_segments),
                    "display_tagged": segments_to_tagged(display_segments),
                    "segments": display_segments,
                    "is_first": (i == 0),
                    "owner_idx": owner_idx if i == 0 else None,
                })
            owner_idx += 1

    return result


def build_visual_rows(content_text: str, owner_text: str = ""):
    """Word 輸出用：先還原成原始段落，再依既有 Word 規則展開成列。"""
    normalized = normalize_content_for_editor(content_text)
    expanded = expand_content_lines(normalized)
    owner_src = [] if not owner_text else str(owner_text).replace("\r\n", "\n").replace("\r", "\n").split("\n")
    rows = []
    for row_no, item in enumerate(expanded, start=1):
        owner_raw = owner_src[row_no - 1].strip() if row_no - 1 < len(owner_src) else ""
        owner_segments = split_styled_segments(owner_raw) if owner_raw else (item.get("owner_inline_segments") or [])
        rows.append({
            "content": item.get("display", ""),
            "content_segments": item.get("segments", []),
            "owner": segments_to_plain(owner_segments) if owner_segments else item.get("owner_inline", ""),
            "owner_segments": owner_segments,
        })
    if not rows:
        rows = [{"content": "", "content_segments": [], "owner": "", "owner_segments": []}]
    return rows


def html_escape_text(s: str) -> str:
    return ("" if s is None else str(s)).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def image_file_to_data_uri(file_path: str):
    try:
        with open(file_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("ascii")
        ext = os.path.splitext(str(file_path))[1].lower()
        mime = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".gif": "image/gif",
            ".bmp": "image/bmp",
            ".webp": "image/webp",
        }.get(ext, "image/png")
        return f"data:{mime};base64,{b64}"
    except Exception:
        return None


def uploaded_file_bytes(uploaded_file):
    if uploaded_file is None:
        return b""
    try:
        if hasattr(uploaded_file, "seek"):
            uploaded_file.seek(0)
        if hasattr(uploaded_file, "getvalue"):
            data = uploaded_file.getvalue()
            if data:
                return data
        if hasattr(uploaded_file, "read"):
            if hasattr(uploaded_file, "seek"):
                uploaded_file.seek(0)
            return uploaded_file.read()
    except Exception:
        return b""
    return b""


def save_uploaded_file_to_temp(uploaded_file):
    data = uploaded_file_bytes(uploaded_file)
    if not data:
        raise ValueError(f"無法讀取上傳檔案：{getattr(uploaded_file, 'name', '未命名')}")
    suffix = os.path.splitext(getattr(uploaded_file, "name", ""))[1] or ".bin"
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.write(data)
    tmp.close()
    return tmp.name


def clear_attachment_upload_widgets():
    for key in [
        "img_upload", "xlsx_upload", "tbl_editor", "tbl_editor_shape",
        "xlsx_preview_idx", "xlsx_sheet"
    ]:
        if key in st.session_state:
            del st.session_state[key]


def table_data_to_preview_html(table_data, caption=""):
    caption_html = ""
    if caption:
        caption_html = f"<div class='preview-att-caption'>{html_escape_text(caption).replace(' ', '&nbsp;')}</div>"

    if not table_data:
        return caption_html + "<div class='preview-table-placeholder'>📊&nbsp;空白表格</div>"

    rows_html = []
    max_cols = max(len(r) for r in table_data) if table_data else 0
    for r_idx, row in enumerate(table_data):
        cells = []
        for c_idx in range(max_cols):
            val = row[c_idx] if c_idx < len(row) else ""
            cell_html = segments_to_html(split_styled_segments(str(val))) if str(val) else "&nbsp;"
            tag = "th" if r_idx == 0 else "td"
            cells.append(f"<{tag}>{cell_html}</{tag}>")
        rows_html.append(f"<tr>{''.join(cells)}</tr>")

    return (
        f"<div class='preview-att-wrap'>"
        f"{caption_html}"
        f"<div class='preview-table-wrap'><table class='preview-inner-table'>{''.join(rows_html)}</table></div>"
        f"</div>"
    )


def render_word_like_preview(content_text: str, owner_text: str, attachments=None):
    attachments = attachments or []
    normalized = normalize_content_for_editor(content_text)
    expanded = expand_content_lines(normalized)
    owner_src = [] if not owner_text else str(owner_text).replace("\r\n", "\n").replace("\r", "\n").split("\n")

    def owner_html_for(item, row_no):
        owner_raw = owner_src[row_no - 1].strip() if row_no - 1 < len(owner_src) else ""
        if owner_raw:
            segs = split_styled_segments(owner_raw)
        else:
            segs = item.get("owner_inline_segments") or split_styled_segments(item.get("owner_inline", ""))
        return segments_to_html(segs) if segs else "<span class='ghost-line'>·&nbsp;·&nbsp;·</span>"

    def build_attachment_block(att):
        row_cost = max(1, int(att.get("row_cost", 1)))
        caption = html_escape_text(att.get("caption", "")).replace(" ", "&nbsp;")

        if att.get("type") == "image":
            uri = image_file_to_data_uri(att.get("path", "")) if att.get("path") else None
            if uri:
                img_html = f"<img src='{uri}' class='preview-img-single'>"
            else:
                img_html = "<div class='preview-img-missing'>圖片載入失敗</div>"
            cap_top = f"<div class='preview-att-caption'>{caption}</div>" if caption and att.get("caption_pos") == "above" else ""
            cap_bottom = f"<div class='preview-att-caption'>{caption}</div>" if caption and att.get("caption_pos") == "below" else ""
            preview_block = f"<div class='preview-att-wrap'>{cap_top}<div class='preview-img-box single'>{img_html}</div>{cap_bottom}</div>"

        elif att.get("type") == "image_pair":
            left_uri = image_file_to_data_uri(att.get("left_path", "")) if att.get("left_path") else None
            right_uri = image_file_to_data_uri(att.get("right_path", "")) if att.get("right_path") else None
            left_html = f"<img src='{left_uri}' class='preview-img-pair'>" if left_uri else "<div class='preview-img-missing'>左圖失敗</div>"
            right_html = f"<img src='{right_uri}' class='preview-img-pair'>" if right_uri else "<div class='preview-img-missing'>右圖失敗</div>"
            pair_top = f"<div class='preview-att-caption'>{caption}</div>" if caption and att.get("caption_pos") == "above" else ""
            pair_bottom = f"<div class='preview-att-caption'>{caption}</div>" if caption and att.get("caption_pos") == "below" else ""
            left_cap = html_escape_text(att.get("left_caption", "")).replace(" ", "&nbsp;")
            right_cap = html_escape_text(att.get("right_caption", "")).replace(" ", "&nbsp;")
            preview_block = (
                f"<div class='preview-att-wrap'>{pair_top}"
                f"<div class='preview-img-box pair'>"
                f"<div class='pair-item'><div class='pair-cap'>{left_cap}</div>{left_html}</div>"
                f"<div class='pair-item'><div class='pair-cap'>{right_cap}</div>{right_html}</div>"
                f"</div>{pair_bottom}</div>"
            )

        elif att.get("type") == "table":
            preview_block = table_data_to_preview_html(att.get("data", []), att.get("caption", ""))

        else:
            label = html_escape_text(att.get("label", "附件"))
            preview_block = f"<div class='preview-table-placeholder'>📎&nbsp;{label}</div>"

        return {"kind": "attachment", "row_cost": row_cost, "block": preview_block}

    items = []
    att_map = {}
    for att in attachments:
        att_map.setdefault(int(att.get("after_line", 0)), []).append(att)

    for att in att_map.get(0, []):
        items.append(build_attachment_block(att))

    if expanded:
        for row_no, item in enumerate(expanded, start=1):
            left_html = segments_to_html(item.get("segments") or []) if item.get("segments") else "<span class='ghost-line'>·&nbsp;·&nbsp;·&nbsp;·&nbsp;·</span>"
            right_html = owner_html_for(item, row_no)
            items.append({"kind": "text", "row_cost": 1, "left_html": left_html, "right_html": right_html})
            for att in att_map.get(row_no, []):
                items.append(build_attachment_block(att))
    else:
        items.append({"kind": "text", "row_cost": 1, "left_html": "<span class='ghost-line'>·&nbsp;·&nbsp;·&nbsp;·&nbsp;·</span>", "right_html": "<span class='ghost-line'>·&nbsp;·&nbsp;·</span>"})
        for att in att_map.get(0, []):
            items.append(build_attachment_block(att))

    def paginate_preview_items(items):
        if not items:
            return [[{"kind": "text", "row_cost": 1, "left_html": "<span class='ghost-line'>·&nbsp;·&nbsp;·&nbsp;·&nbsp;·</span>", "right_html": "<span class='ghost-line'>·&nbsp;·&nbsp;·</span>"}]]
        pages = []
        current = []
        page_no = 1
        capacity = 11
        used = 0
        for item in items:
            cost = max(1, int(item.get("row_cost", 1)))
            page_capacity = 11 if page_no == 1 else 26
            if current and used + cost > page_capacity:
                pages.append(current)
                current = []
                page_no += 1
                used = 0
                page_capacity = 11 if page_no == 1 else 26
            current.append(item)
            used += cost
        if current:
            pages.append(current)
        return pages

    def render_page_rows(page_items):
        rows = []
        for item in page_items:
            if item["kind"] == "text":
                rows.append(f"<tr><td class='preview-left'>{item['left_html']}</td><td class='preview-right'>{item['right_html']}</td></tr>")
            else:
                row_cost = max(1, int(item.get("row_cost", 1)))
                rows.append(f"<tr><td class='preview-left attachment-cell' rowspan='{row_cost}'>{item['block']}</td><td class='preview-right attachment-cell' rowspan='{row_cost}'></td></tr>")
                for _ in range(row_cost - 1):
                    rows.append("<tr></tr>")
        return ''.join(rows)

    pages = paginate_preview_items(items)
    total_row_count = sum(max(1, int(i.get("row_cost", 1))) for i in items)
    total_pages = len(pages)

    page_blocks = []
    for idx, page_items in enumerate(pages, start=1):
        display = "block" if idx == 1 else "none"
        page_blocks.append(f"""
        <div class='word-preview-paper' id='preview-page-{idx}' style='display:{display}'>
          <div class='page-watermark'>第 {idx} 頁 / 共 {total_pages} 頁</div>
          <div class='word-preview-title'>👁️ 表格預覽（依 Word 帶入邏輯）</div>
          <table class='word-preview-table'>
            <thead>
              <tr>
                <th class='preview-left'>決　議　事　項</th>
                <th class='preview-right'>決議執行者</th>
              </tr>
            </thead>
            <tbody>
              {render_page_rows(page_items)}
            </tbody>
          </table>
        </div>
        """)

    nav_bar = f"""
    <div class='preview-nav'>
      <button class='nav-btn' id='btn-prev' onclick='changePage(-1)' disabled>&#9664; 上一頁</button>
      <span class='nav-info' id='nav-info'>第 1 頁 / 共 {total_pages} 頁</span>
      <button class='nav-btn' id='btn-next' onclick='changePage(1)' {'disabled' if total_pages <= 1 else ''}>下一頁 &#9654;</button>
    </div>
    """ if total_pages >= 1 else ""

    preview_html = f"""
    <style>
    body {{ margin:0; padding:0; background:#f3f3f3; }}
    .preview-nav {{
        display:flex; align-items:center; justify-content:center;
        gap:16px; padding:8px 0 10px; background:#f3f3f3;
        position:sticky; top:0; z-index:99;
        border-bottom:1px solid #ddd; margin-bottom:10px;
    }}
    .nav-btn {{
        background:#5b4a36; color:#fff; border:none; border-radius:6px;
        padding:5px 18px; font-size:12pt; cursor:pointer;
        font-family:'Microsoft JhengHei','PingFang TC',sans-serif;
    }}
    .nav-btn:disabled {{ background:#bbb; cursor:default; }}
    .nav-info {{
        font-size:12pt; min-width:130px; text-align:center;
        font-family:'Microsoft JhengHei','PingFang TC',sans-serif; color:#333;
    }}
    .word-preview-wrap {{ background:#f3f3f3; padding:4px 8px 20px; }}
    .word-preview-paper {{
        width:min(21cm, calc(100vw - 32px)); min-height:29.7cm;
        margin:0 auto 22px auto; box-sizing:border-box;
        background:#ffffff; padding:2cm 1.5cm 2cm 2cm;
        box-shadow:0 0 0 1px rgba(0,0,0,.10), 0 6px 20px rgba(0,0,0,.08);
        position:relative;
    }}
    .page-watermark {{
        position:absolute; top:10px; right:18px; font-size:11pt; color:#999;
        font-family:'Microsoft JhengHei','PingFang TC',sans-serif;
    }}
    .word-preview-title {{
        margin:0 0 8px 0; font-size:13.5pt; font-weight:700; color:#222222;
        font-family:'Microsoft JhengHei','PingFang TC',sans-serif;
    }}
    table.word-preview-table {{
        width:100%; border-collapse:collapse; table-layout:fixed;
        font-family:'DFKai-SB','BiauKai','KaiTi','標楷體',serif;
        font-size:13.5pt; color:#111111; background:#fff; border:1.6px solid #000;
    }}
    table.word-preview-table th,
    table.word-preview-table td {{
        border:1px solid #000; padding:2px 6px; vertical-align:middle;
        line-height:1.35; white-space:pre-wrap; word-break:break-word; overflow-wrap:break-word;
    }}
    table.word-preview-table th {{
        text-align:center; font-weight:700; background:#fff;
        letter-spacing:.18em; height:34px; font-size:13.5pt;
    }}
    .preview-left {{
        width:85.8%; min-height:28px; height:28px; text-align:left;
        background-image:radial-gradient(circle,rgba(0,0,0,.22) 0.48px,transparent 0.62px);
        background-size:9px 9px; background-position:center center;
    }}
    .preview-right {{
        width:14.2%; text-align:center; min-height:28px; height:28px;
        background-image:radial-gradient(circle,rgba(0,0,0,.18) 0.48px,transparent 0.62px);
        background-size:9px 9px; background-position:center center;
    }}
    .attachment-cell {{ background-image:none; background:#fff; padding:6px 8px; vertical-align:top; }}
    .preview-att-wrap {{ display:flex; flex-direction:column; gap:6px; align-items:center; }}
    .preview-att-caption {{ font-size:12pt; font-weight:700; text-align:center; width:100%; }}
    .preview-img-box.single {{ width:100%; display:flex; justify-content:center; }}
    .preview-img-box.pair {{ width:100%; display:grid; grid-template-columns:1fr 1fr; gap:10px; align-items:start; }}
    .pair-item {{ display:flex; flex-direction:column; align-items:center; gap:4px; }}
    .pair-cap {{ min-height:1.2em; font-size:11pt; font-weight:700; text-align:center; }}
    .preview-img-single,.preview-img-pair {{ max-width:100%; max-height:220px; object-fit:contain; border:1px solid #888; }}
    .preview-img-missing,.preview-table-placeholder {{ color:#666; font-size:12pt; text-align:center; padding:18px 8px; border:1px dashed #aaa; width:100%; box-sizing:border-box; }}
    .preview-table-wrap {{ width:100%; overflow:auto; }}
    .preview-inner-table {{ width:100%; border-collapse:collapse; table-layout:fixed; background:#fff; font-size:12pt; }}
    .preview-inner-table th,.preview-inner-table td {{ border:1px solid #666; padding:2px 4px; line-height:1.35; vertical-align:middle; white-space:pre-wrap; word-break:break-word; overflow-wrap:break-word; background:#fff; }}
    .preview-inner-table th {{ font-weight:700; text-align:center; }}
    .ghost-line {{ color:#c3c3c3; letter-spacing:.08em; }}
    </style>
    <script>
    var _curPage = 1;
    var _totalPages = {total_pages};
    function changePage(delta) {{
        var next = _curPage + delta;
        if(next < 1 || next > _totalPages) return;
        document.getElementById('preview-page-' + _curPage).style.display = 'none';
        _curPage = next;
        document.getElementById('preview-page-' + _curPage).style.display = 'block';
        document.getElementById('nav-info').textContent = '第 ' + _curPage + ' 頁 / 共 ' + _totalPages + ' 頁';
        document.getElementById('btn-prev').disabled = (_curPage === 1);
        document.getElementById('btn-next').disabled = (_curPage === _totalPages);
        window.scrollTo(0, 0);
    }}
    </script>
    {nav_bar}
    <div class='word-preview-wrap'>
      {''.join(page_blocks)}
    </div>
    """
    est_height = 220 + len(pages) * 760 + total_row_count * 4
    st.components.v1.html(preview_html, height=min(3200, max(420, est_height)), scrolling=True)

def normalize_content_for_editor(text: str) -> str:
    """
    將 editor 內顯示文字還原成「原始段落」：
    - 一般首行：INDENT + 內容
    - 自動換行接續行：以 CONT_MARK + 內容 標記（零寬字元，不可見）
    - 使用者手動按 Enter 另起的新行：不帶 CONT_MARK，視為新段落
    - 決議事項行：原樣保留
    - 附件佔位符行：跳過

    這樣可解決：按過一次 Word排版後，再手動 Enter 新增文字，
    第二次按 Word排版時，新的段落不會被誤併回上一段未滿 60 字元的內容。
    """
    if text is None:
        return ""

    lines = str(text).replace("\r\n", "\n").replace("\r", "\n").split("\n")
    raw_lines = []

    for line in lines:
        if line.strip().startswith(_ATT_PREFIX) and line.strip().endswith(_ATT_SUFFIX):
            continue

        if not line.strip():
            raw_lines.append("")
            continue

        if split_content_owner_line(line) is not None:
            raw_lines.append(line)
            continue

        if line.startswith(CONT_MARK):
            line = line[len(CONT_MARK):]
            if raw_lines:
                line = strip_continuation_display_indent(raw_lines[-1], line)
                raw_lines[-1] = raw_lines[-1] + line
            else:
                raw_lines.append(line)
            continue

        stripped_line = strip_display_indent(line)
        if stripped_line != line:
            raw_lines.append(stripped_line)
            continue

        # 沒有縮排、也沒有 CONT_MARK：視為使用者手動 Enter 的新段落
        raw_lines.append(line)

    return "\n".join(raw_lines)




def build_row_map_cache(content_text: str):
    """
    以目前 editor 內容為基礎，先 normalize 成原始段落，再依 Word 規則建立列映射快取。
    之後搜尋、列數統計、附件插入位置都直接讀這個快取。
    """
    visual_rows = []
    paragraph_map = []

    normalized = normalize_content_for_editor(content_text)
    raw_lines = [] if not normalized else str(normalized).replace("\r\n", "\n").replace("\r", "\n").split("\n")
    current_visual = 1

    for raw_idx, raw in enumerate(raw_lines, start=1):
        start_row = current_visual

        if not raw.strip():
            visual_rows.append({
                "visual_row": current_visual,
                "raw_row": raw_idx,
                "content": "",
                "is_resolution": False,
            })
            current_visual += 1
        else:
            parsed = split_content_owner_line(raw)
            if parsed:
                visual_rows.append({
                    "visual_row": current_visual,
                    "raw_row": raw_idx,
                    "content": raw,
                    "is_resolution": True,
                })
                current_visual += 1
            else:
                first_indent = get_display_indent(raw)
                first_indent_units = detect_outline_indent_units(raw)
                segs = wrap_plain_line_30(raw, initial_units=first_indent_units)
                for i, seg in enumerate(segs):
                    display = first_indent + seg if i == 0 else seg
                    visual_rows.append({
                        "visual_row": current_visual,
                        "raw_row": raw_idx,
                        "content": display,
                        "is_resolution": False,
                    })
                    current_visual += 1

        end_row = current_visual - 1
        paragraph_map.append({
            "raw_row": raw_idx,
            "start_visual_row": start_row,
            "end_visual_row": end_row,
            "source": raw,
        })

    return {
        "visual_rows": visual_rows,
        "paragraph_map": paragraph_map,
        "visual_count": len(visual_rows),
        "raw_count": len(raw_lines),
    }


def refresh_row_map_cache():
    content = st.session_state.get("content", "")
    cache = build_row_map_cache(content)
    st.session_state["row_map_cache"] = cache
    st.session_state["_row_map_hash"] = hash(content)
    return cache


def get_row_map_cache():
    content = st.session_state.get("content", "")
    content_hash = hash(content)
    cache = st.session_state.get("row_map_cache")

    if cache is None or st.session_state.get("_row_map_hash") != content_hash:
        cache = build_row_map_cache(content)
        st.session_state["row_map_cache"] = cache
        st.session_state["_row_map_hash"] = content_hash

    return cache


def build_editor_display(content_text: str, attachments: list) -> str:
    """
    給 editor 顯示的字串：目前純文字內容 + 附件佔位符。
    不在輸入當下自動展開/排版，避免游標跳動與閃爍。
    """
    import collections

    src_lines = [] if not content_text else str(content_text).replace("\r\n", "\n").replace("\r", "\n").split("\n")

    att_after: dict = collections.defaultdict(list)
    for att in sorted(attachments, key=lambda a: int(a["after_line"])):
        label = att.get("label", "附件")
        cost = int(att.get("row_cost", 1))
        att_after[int(att["after_line"])].extend([make_placeholder(label)] * cost)

    result = list(att_after.get(0, []))
    visual_row = 0

    for line in src_lines:
        result.append(line)
        visual_row += 1
        result.extend(att_after.get(visual_row, []))

    if not src_lines and att_after.get(0):
        return "\n".join(result)
    return "\n".join(result)


def format_content_for_word_button(content_text: str) -> str:
    """
    手動按下「Word排版」時使用：
    1. 先把目前 editor 內容 normalize 成原始段落
    2. 再依 Word 規則重新展開
    3. 對「自動換行產生的接續行」加上 CONT_MARK（零寬字元）

    之後若使用者在 editor 手動按 Enter 另起新行，該新行不會帶 CONT_MARK，
    下一次 Word排版時就能正確保留為新段落，而不會被誤補回前一行。
    """
    normalized = normalize_content_for_editor(content_text)
    expanded = expand_content_lines(normalized)
    rendered = []
    for item in expanded:
        display = item["display"]
        if not item.get("is_first", True) and not item.get("is_resolution"):
            display = CONT_MARK + display
        rendered.append(display)
    return "\n".join(rendered)



def normalize_owner_text(text: str, target_rows: int):
    lines = [] if text is None else str(text).replace("\r\n", "\n").replace("\r", "\n").split("\n")
    if len(lines) < target_rows:
        lines += [""] * (target_rows - len(lines))
    else:
        lines = lines[:target_rows]
    return "\n".join(lines)


def owner_lines_from_text(text: str, target_rows: int):
    return normalize_owner_text(text, target_rows).split("\n") if target_rows > 0 else []


def sync_owner_rows():
    content = st.session_state.get("content", "") or ""
    normalized = normalize_content_for_editor(content)
    expanded = expand_content_lines(normalized)
    # owner 列數 = 目前會議記錄顯示列數；使用者打在哪一列，Word 就帶哪一列
    target_rows = len(expanded)
    new_owner = normalize_owner_text(st.session_state.get("owner_text", ""), target_rows)
    if new_owner != st.session_state.get("owner_text", ""):
        st.session_state["owner_text"] = new_owner


def calc_total_rows_from_content(text: str, attachments: list):
    rows = build_visual_rows(text, st.session_state.get("owner_text", ""))
    text_rows = len(rows)
    att_rows = sum(int(a.get("row_cost", 1)) for a in attachments)
    return text_rows + att_rows, text_rows, att_rows


def calc_page_stats(content_text: str, attachments: list):
    total_rows, text_rows, att_rows = calc_total_rows_from_content(content_text, attachments)
    if total_rows <= P1_ROWS:
        page_num = 1
        rows_in_page = total_rows
    else:
        overflow = total_rows - P1_ROWS
        page_num = 2 + (overflow - 1) // P2_ROWS
        rows_in_page = ((overflow - 1) % P2_ROWS) + 1
    return {
        "total_rows": total_rows,
        "text_rows": text_rows,
        "att_rows": att_rows,
        "page_num": page_num,
        "rows_in_page": rows_in_page,
    }


def calc_total_pages(content_text: str, attachments: list):
    return calc_page_stats(content_text, attachments)["page_num"]


def set_cell_text(cell, text):
    cell.text = str(text) if text is not None else ""
    for p in cell.paragraphs:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            style_run_kaiti(run, size=14)


def style_run_kaiti(run, size=10, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "標楷體"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "標楷體")
    if color:
        run.font.color.rgb = color


def set_cell_rich_text(cell, segments, align=WD_ALIGN_PARAGRAPH.LEFT, size=14):
    cell.text = ""
    if not cell.paragraphs:
        cell.add_paragraph("")
    p = cell.paragraphs[0]
    p.paragraph_format.alignment = align
    if not segments:
        return
    for txt, color_key in segments:
        if txt == "":
            continue
        run = p.add_run(txt)
        color = None
        if color_key:
            try:
                from docx.shared import RGBColor
                color = RGBColor.from_string(COLOR_NAME_TO_HEX.get(color_key, "000000"))
            except Exception:
                color = None
        style_run_kaiti(run, size=size, color=color)


def get_distinct_cells(row):
    distinct = []
    for c in row.cells:
        if not distinct or c._tc != distinct[-1]._tc:
            distinct.append(c)
    return distinct


def remove_page_breaks_in_paragraph(p):
    for br in p._element.xpath(".//w:br"):
        try:
            if br.get(qn("w:type")) == "page":
                br.getparent().remove(br)
        except Exception:
            pass
    for lr in p._element.xpath(".//w:lastRenderedPageBreak"):
        try:
            lr.getparent().remove(lr)
        except Exception:
            pass


def cleanup_cell_paragraphs(cell):
    try:
        ps = cell._tc.xpath("./w:p")
        if not ps:
            return
        for p in cell.paragraphs:
            remove_page_breaks_in_paragraph(p)
        for p in ps[1:]:
            p.getparent().remove(p)
        if not cell.paragraphs:
            cell.add_paragraph("")
    except Exception:
        pass


def find_row_index_contains(table, predicate):
    for r_idx, row in enumerate(table.rows):
        clean_text = "".join(c.text for c in get_distinct_cells(row))
        clean_text = clean_text.replace(" ", "").replace("\n", "").replace("　", "")
        if predicate(clean_text):
            return r_idx
    return -1


def _apply_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:color"), "000000")
                tcBorders.append(border)
            tcPr.append(tcBorders)


def _clear_cell(cell):
    cell.text = ""
    for p in cell.paragraphs:
        try:
            p.clear()
        except Exception:
            pass


def _make_caption_para(cell, caption):
    p = cell.add_paragraph() if (cell.paragraphs and (cell.paragraphs[0].text or cell.paragraphs[0].runs)) else cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    style_run_kaiti(run, size=11, bold=True)


def insert_image_into_row(row, img_path, row_cost, caption="", caption_pos="above"):
    dist_cells = get_distinct_cells(row)
    cell = dist_cells[0]
    _clear_cell(cell)
    img_height_cm = max(1.0, (row_cost - (1 if caption else 0)) * 0.7)

    def _make_img_para(c):
        p = c.add_paragraph() if c.paragraphs[0].runs or c.paragraphs[0].text else c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(img_path, height=Cm(img_height_cm))
        except Exception:
            run.text = f"[圖片載入失敗: {os.path.basename(img_path)}]"
            style_run_kaiti(run, size=10)

    if caption and caption_pos == "above":
        _make_caption_para(cell, caption)
        _make_img_para(cell)
    elif caption and caption_pos == "below":
        _make_img_para(cell)
        _make_caption_para(cell, caption)
    else:
        _make_img_para(cell)

    if len(dist_cells) > 1:
        set_cell_text(dist_cells[-1], "")


def insert_image_pair_into_row(row, left_path, right_path, row_cost, caption="", caption_pos="above", left_caption="", right_caption=""):
    dist_cells = get_distinct_cells(row)
    cell = dist_cells[0]
    _clear_cell(cell)
    extra_caption_rows = 1 if caption else 0
    extra_inner_rows = 1 if (left_caption or right_caption) else 0
    img_height_cm = max(1.0, (row_cost - extra_caption_rows - extra_inner_rows) * 0.62)

    if caption and caption_pos == "above":
        _make_caption_para(cell, caption)

    nested_rows = 2 if (left_caption or right_caption) else 1
    nested = cell.add_table(rows=nested_rows, cols=2)
    try:
        nested.style = "Table Grid"
    except Exception:
        pass

    if nested_rows == 2:
        for idx, cap in enumerate([left_caption, right_caption]):
            cp = nested.rows[0].cells[idx].paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cap:
                run = cp.add_run(cap)
                style_run_kaiti(run, size=11, bold=True)

    img_row_idx = 1 if nested_rows == 2 else 0
    for idx, img_path in enumerate([left_path, right_path]):
        c = nested.rows[img_row_idx].cells[idx]
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            if img_path:
                run.add_picture(img_path, width=Cm(6.0), height=Cm(img_height_cm))
            else:
                run.text = ""
        except Exception:
            run.text = f"[圖片載入失敗: {os.path.basename(img_path)}]"
            style_run_kaiti(run, size=10)

    if caption and caption_pos == "below":
        _make_caption_para(cell, caption)

    if len(dist_cells) > 1:
        set_cell_text(dist_cells[-1], "")


def insert_table_into_row(row, table_data, caption="", caption_pos="above"):
    dist_cells = get_distinct_cells(row)
    cell = dist_cells[0]
    for p in cell.paragraphs:
        p.clear()

    def _add_caption(c):
        p = c.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        style_run_kaiti(run, size=11, bold=True)

    if not table_data or not table_data[0]:
        cell.paragraphs[0].add_run("[空白表格]")
        return

    if caption and caption_pos == "above":
        cap_p = cell.paragraphs[0]
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(caption)
        style_run_kaiti(cap_run, size=11, bold=True)

    nested = cell.add_table(rows=len(table_data), cols=len(table_data[0]))
    for _style_name in ("格線表格", "Table Grid", "TableGrid"):
        try:
            nested.style = _style_name
            break
        except KeyError:
            continue
    else:
        _apply_table_borders(nested)

    for r_i, row_data in enumerate(table_data):
        for c_i, val in enumerate(row_data):
            c = nested.rows[r_i].cells[c_i]
            c.text = str(val)
            for p in c.paragraphs:
                for run in p.runs:
                    style_run_kaiti(run, size=11, bold=(r_i == 0))

    if caption and caption_pos == "below":
        _add_caption(cell)

    if len(dist_cells) > 1:
        set_cell_text(dist_cells[-1], "")


def _merge_rows_left_cell(table, start_row, end_row):
    if end_row <= start_row or end_row >= len(table.rows):
        return
    try:
        start_cell = get_distinct_cells(table.rows[start_row])[0]
        end_cell = get_distinct_cells(table.rows[end_row])[0]
        merged = start_cell.merge(end_cell)
        cleanup_cell_paragraphs(merged)
    except Exception:
        pass


def attachment_to_chunk(att):
    return {
        "type": att["type"],
        "data": att.get("data"),
        "path": att.get("path"),
        "left_path": att.get("left_path"),
        "right_path": att.get("right_path"),
        "row_cost": int(att["row_cost"]),
        "caption": att.get("caption", ""),
        "left_caption": att.get("left_caption", ""),
        "right_caption": att.get("right_caption", ""),
        "caption_pos": att.get("caption_pos", "above"),
    }


def build_chunks(content_text, owner_text, attachments):
    rows = build_visual_rows(content_text, owner_text)

    text_chunks = []
    for line_no, row in enumerate(rows, start=1):
        text_chunks.append({
            "type": "text",
            "content": row["content"],
            "content_segments": row.get("content_segments", []),
            "owner": row["owner"],
            "owner_segments": row.get("owner_segments", []),
            "line_no": line_no,
        })

    sorted_atts = sorted(attachments, key=lambda a: int(a["after_line"]))
    result = []
    att_idx = 0

    while att_idx < len(sorted_atts) and int(sorted_atts[att_idx]["after_line"]) == 0:
        att = sorted_atts[att_idx]
        result.append(attachment_to_chunk(att))
        att_idx += 1

    current_line = 0
    for chunk in text_chunks:
        result.append(chunk)
        if chunk["line_no"] != current_line:
            current_line = chunk["line_no"]
            while att_idx < len(sorted_atts) and int(sorted_atts[att_idx]["after_line"]) == current_line:
                att = sorted_atts[att_idx]
                result.append(attachment_to_chunk(att))
                att_idx += 1

    while att_idx < len(sorted_atts):
        att = sorted_atts[att_idx]
        result.append(attachment_to_chunk(att))
        att_idx += 1

    return result


def build_chunks_for_word(content_text, owner_text, attachments):
    """
    Word 產出專用版本：
    - 一個視覺列 = 一個文字 chunk，直接逐列表格寫入
    - after_line 以視覺列號為準，可與畫面列數一致
    """
    rows = build_visual_rows(content_text, owner_text)
    text_chunks = []
    for line_no, row in enumerate(rows, start=1):
        text_chunks.append({
            "type": "text",
            "content": row["content"],
            "content_segments": row.get("content_segments", []),
            "owner": row["owner"],
            "owner_segments": row.get("owner_segments", []),
            "line_no": line_no,
            "visual_cost": 1,
        })

    sorted_atts = sorted(attachments, key=lambda a: int(a["after_line"]))
    result = []
    att_idx = 0

    while att_idx < len(sorted_atts) and int(sorted_atts[att_idx]["after_line"]) == 0:
        att = sorted_atts[att_idx]
        result.append(attachment_to_chunk(att))
        att_idx += 1

    current_line = 0
    for chunk in text_chunks:
        result.append(chunk)
        current_line = chunk["line_no"]
        while att_idx < len(sorted_atts) and int(sorted_atts[att_idx]["after_line"]) == current_line:
            att = sorted_atts[att_idx]
            result.append(attachment_to_chunk(att))
            att_idx += 1

    while att_idx < len(sorted_atts):
        att = sorted_atts[att_idx]
        result.append(attachment_to_chunk(att))
        att_idx += 1

    return result

def chunk_row_cost(chunk):
    if chunk["type"] == "text":
        return chunk.get("visual_cost", 1)
    return int(chunk.get("row_cost", 6))


def fill_chunk_into_row(row, chunk):
    if chunk["type"] == "text":
        dist_cells = get_distinct_cells(row)
        set_cell_rich_text(dist_cells[0], chunk.get("content_segments", []), align=WD_ALIGN_PARAGRAPH.LEFT, size=14)
        if len(dist_cells) > 1:
            set_cell_rich_text(dist_cells[-1], chunk.get("owner_segments", []), align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    elif chunk["type"] == "image":
        insert_image_into_row(
            row,
            chunk["path"],
            int(chunk.get("row_cost", 6)),
            chunk.get("caption", ""),
            chunk.get("caption_pos", "above"),
        )
    elif chunk["type"] == "image_pair":
        insert_image_pair_into_row(
            row,
            chunk.get("left_path"),
            chunk.get("right_path"),
            int(chunk.get("row_cost", 6)),
            chunk.get("caption", ""),
            chunk.get("caption_pos", "above"),
            chunk.get("left_caption", ""),
            chunk.get("right_caption", ""),
        )
    elif chunk["type"] == "table":
        insert_table_into_row(
            row,
            chunk["data"],
            chunk.get("caption", ""),
            chunk.get("caption_pos", "above"),
        )


def generate_doc_bytes(data: dict, content_text: str, owner_text: str, attachments: list) -> io.BytesIO:
    template_name = "總經理範本.docx" if data["highest"] == "總經理" else "副總經理範本.docx"
    template_path = os.path.join(TEMPLATE_DIR, template_name)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"找不到範本：{template_path}\n請確認 templates/ 資料夾內有 {template_name}")

    doc = Document(template_path)
    if not doc.tables:
        raise ValueError("範本內找不到表格")

    tables = doc.tables
    table1 = tables[0]

    search_targets = {
        "title": lambda t: "會議名稱" in t,
        "time": lambda t: "開會時間" in t or "會議時間" in t,
        "location": lambda t: "開會地點" in t,
        "chair": lambda t: "主席" in t,
        "recorder": lambda t: "會議紀錄" in t or "會議記錄" in t,
        "attendees": lambda t: ("出" in t and "席" in t and "員" in t),
    }

    for row in table1.rows:
        dist_cells = get_distinct_cells(row)
        for i, cell in enumerate(dist_cells):
            clean_text = cell.text.replace(" ", "").replace("\n", "").replace("　", "")
            for key, condition in list(search_targets.items()):
                if condition(clean_text) and i + 1 < len(dist_cells):
                    set_cell_text(dist_cells[i + 1], data[key])
                    del search_targets[key]
                    break

    chunks = build_chunks_for_word(content_text, owner_text, attachments)

    resolution_start = find_row_index_contains(table1, lambda t: ("決議事項" in t) and ("執行" in t or "決議執行者" in t))
    if resolution_start == -1:
        raise ValueError("找不到第一頁『決議事項/決議執行者』標題列")

    signature_start = find_row_index_contains(
        table1,
        lambda t: ("總經理" in t) or ("副總經理" in t) or ("助理副總" in t) or ("承辦單位" in t)
    )
    if signature_start == -1:
        raise ValueError("找不到簽核欄")
    if signature_start <= resolution_start:
        raise ValueError("簽核欄位置在決議標題之前，請檢查範本")

    available_slots_page1 = min(signature_start - resolution_start - 1, P1_ROWS)
    if available_slots_page1 < 1:
        raise ValueError("第一頁決議可用列數為 0")

    used = 0
    page1_end = 0
    for i, chunk in enumerate(chunks):
        cost = chunk_row_cost(chunk)
        if used + cost > available_slots_page1:
            break
        used += cost
        page1_end = i + 1

    page1_chunks = chunks[:page1_end]
    leftover_chunks = chunks[page1_end:]

    for r in range(resolution_start + 1, signature_start):
        dist_cells = get_distinct_cells(table1.rows[r])
        for c in dist_cells:
            set_cell_text(c, "")
            for p in c.paragraphs:
                remove_page_breaks_in_paragraph(p)

    slot_idx = resolution_start + 1
    for chunk in page1_chunks:
        cost = chunk_row_cost(chunk)
        if chunk["type"] != "text" and cost > 1:
            # 附件才 merge（圖片/表格需要連續空間）
            _merge_rows_left_cell(table1, slot_idx, min(slot_idx + cost - 1, signature_start - 1))
        fill_chunk_into_row(table1.rows[slot_idx], chunk)
        # 文字 chunk 佔多列時，後續列留空（格線保留，不 merge）
        if chunk["type"] == "text" and cost > 1:
            for extra in range(1, cost):
                if slot_idx + extra < signature_start:
                    dist = get_distinct_cells(table1.rows[slot_idx + extra])
                    for c in dist:
                        set_cell_text(c, "")
        slot_idx += cost

    used_table_count = 1
    if leftover_chunks:
        if len(tables) < 2:
            raise ValueError("範本沒有第 2 頁表格")

        idx = 0
        total = len(leftover_chunks)

        for t_idx in range(1, len(tables)):
            tbl = tables[t_idx]
            header_txt = "".join(c.text for c in get_distinct_cells(tbl.rows[0]))
            header_txt = header_txt.replace(" ", "").replace("\n", "").replace("　", "")
            if ("決議事項" not in header_txt) or ("執行" not in header_txt and "決議執行者" not in header_txt):
                raise ValueError(f"第 {t_idx + 1} 張表格第 1 列不是『決議事項/決議執行者』表頭，請檢查範本。")

            data_rows = len(tbl.rows) - 1
            if data_rows <= 0:
                continue

            for r in range(1, len(tbl.rows)):
                dist_cells = get_distinct_cells(tbl.rows[r])
                for c in dist_cells:
                    set_cell_text(c, "")
                    cleanup_cell_paragraphs(c)

            used_rows = 0
            row_pos = 1
            page_has_any = False

            while idx < total:
                chunk = leftover_chunks[idx]
                cost = chunk_row_cost(chunk)
                if used_rows + cost > data_rows:
                    break
                if cost > 1 and chunk["type"] != "text":
                    _merge_rows_left_cell(tbl, row_pos, row_pos + cost - 1)
                fill_chunk_into_row(tbl.rows[row_pos], chunk)
                if chunk["type"] == "text" and cost > 1:
                    for extra in range(1, cost):
                        if row_pos + extra < len(tbl.rows):
                            dist = get_distinct_cells(tbl.rows[row_pos + extra])
                            for c in dist:
                                set_cell_text(c, "")
                used_rows += cost
                row_pos += cost
                idx += 1
                page_has_any = True

            if page_has_any:
                used_table_count = t_idx + 1
            else:
                break

            if idx >= total:
                break

        if idx < total:
            raise ValueError(f"內容太多，放不下（還剩 {total - idx} 列/附件）。請增加範本頁數或縮減內容。")

    for i in range(len(doc.tables) - 1, used_table_count - 1, -1):
        tbl = doc.tables[i]
        tbl._element.getparent().remove(tbl._element)

    if doc.tables:
        last_table = doc.tables[-1]
        body_elements = list(doc.element.body)
        last_tbl_idx = body_elements.index(last_table._element)
        paragraphs_after = [p for p in doc.paragraphs if body_elements.index(p._element) > last_tbl_idx]
        for p in paragraphs_after:
            try:
                p._element.getparent().remove(p._element)
            except Exception:
                pass

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def find_matching_lines(keyword: str):
    cache = get_row_map_cache()
    hits = []
    kw = (keyword or "").strip()
    if not kw:
        return hits

    for row in cache["visual_rows"]:
        if kw in row["content"]:
            hits.append((row["visual_row"], row["content"]))
    return hits


st.set_page_config(page_title="會議紀錄助理", page_icon="📋", layout="wide")

st.markdown("""
<style>
:root{
    --bg:#f6f2eb;
    --panel:#fbf8f3;
    --input:#f6eee2;
    --line:#d9cbb8;
    --line-strong:#c9b394;
    --text:#2d261f;
    --muted:#786c5f;
    --accent:#b58b57;
    --accent-soft:rgba(181,139,87,.16);
    --radius:14px;
    --shadow:0 6px 18px rgba(86,67,44,.05);
}

html, body, [class*="css"]{
    font-size:14pt !important;
}

.stApp{
    background:var(--bg) !important;
    color:var(--text) !important;
}

.block-container{
    max-width:1600px !important;
    padding-top:1.2rem !important;
    padding-bottom:2rem !important;
}

h1{
    color:var(--text) !important;
    font-size:26pt !important;
    font-weight:700 !important;
    margin-bottom:1.1rem !important;
    letter-spacing:.01em;
}

h2{
    color:var(--text) !important;
    font-size:17pt !important;
    font-weight:700 !important;
    margin-bottom:.45rem !important;
}

label[data-testid="stWidgetLabel"] p{
    color:var(--text) !important;
    font-size:13.5pt !important;
    font-weight:600 !important;
}

p, span, div{
    color:var(--text);
}

/* input / textarea */
div[data-baseweb="input"] input,
div[data-baseweb="textarea"] textarea{
    background:var(--input) !important;
    color:var(--text) !important;
    border:1px solid var(--line) !important;
    border-radius:12px !important;
    font-size:14pt !important;
    min-height:46px !important;
    transition:border-color .15s ease, box-shadow .15s ease !important;
}

div[data-baseweb="input"] input:hover,
div[data-baseweb="textarea"] textarea:hover{
    border-color:var(--line-strong) !important;
}

div[data-baseweb="input"] input:focus,
div[data-baseweb="textarea"] textarea:focus{
    border-color:var(--accent) !important;
    box-shadow:0 0 0 4px var(--accent-soft) !important;
}

/* select */
div[data-baseweb="select"] > div{
    background:var(--input) !important;
    color:var(--text) !important;
    border:1px solid var(--line) !important;
    border-radius:12px !important;
    min-height:46px !important;
    font-size:14pt !important;
}

div[data-baseweb="select"] > div:hover{
    border-color:var(--line-strong) !important;
}

div[data-baseweb="select"] span{
    color:var(--text) !important;
    font-size:14pt !important;
}

/* buttons */
.stButton > button,
.stDownloadButton > button{
    border:none !important;
    border-radius:12px !important;
    min-height:44px !important;
    font-size:13.5pt !important;
    font-weight:700 !important;
    background:#b58b57 !important;
    color:#fffaf4 !important;
    box-shadow:0 4px 14px rgba(181,139,87,.18) !important;
    transition:transform .12s ease, box-shadow .15s ease !important;
}

.stButton > button:hover,
.stDownloadButton > button:hover{
    transform:translateY(-1px);
    box-shadow:0 8px 18px rgba(181,139,87,.24) !important;
}

button[kind="secondary"]{
    background:#ece3d6 !important;
    color:var(--text) !important;
    box-shadow:none !important;
}

/* metric */
div[data-testid="metric-container"]{
    background:transparent !important;
    border:none !important;
    box-shadow:none !important;
    padding:.1rem 0 !important;
}

div[data-testid="metric-container"] label,
div[data-testid="metric-container"] [data-testid="stMetricLabel"]{
    color:var(--muted) !important;
    font-size:12.4pt !important;
}

div[data-testid="metric-container"] [data-testid="stMetricValue"]{
    color:var(--text) !important;
    font-size:18pt !important;
    font-weight:700 !important;
}

/* expander */
[data-testid="stExpander"]{
    background:var(--panel) !important;
    border:1px solid var(--line) !important;
    border-radius:14px !important;
    box-shadow:var(--shadow) !important;
}

/* tabs */
button[role="tab"]{
    font-size:12.8pt !important;
    color:var(--muted) !important;
    border-radius:10px 10px 0 0 !important;
    padding:.55rem .8rem !important;
}

button[role="tab"][aria-selected="true"]{
    color:var(--text) !important;
    font-weight:700 !important;
}

/* ace - 改成乾淨淡色，不要厚黑框 */
.ace_editor{
    border:1px solid var(--line) !important;
    border-radius:0 !important;
    box-shadow:none !important;
    overflow:hidden !important;
}

.ace_editor,
.ace_editor .ace_scroller,
.ace_editor .ace_content,
.ace_editor .ace_gutter{
    background:#fbfaf8 !important;
    color:var(--text) !important;
}

.ace_editor .ace_gutter{
    border-right:1px solid #e8ddd0 !important;
    color:#8a7e71 !important;
}

.ace_editor .ace_text-layer,
.ace_editor .ace_line{
    font-size:14pt !important;
    line-height:1.72 !important;
}

.ace_editor .ace_gutter-cell{
    font-size:12.4pt !important;
    line-height:1.72 !important;
}

.ace_marker-layer .ace_active-line{
    background:rgba(181,139,87,.08) !important;
}

.ace_marker-layer .ace_selection{
    background:rgba(181,139,87,.16) !important;
}

/* file uploader 不要黑色大塊 */
[data-testid="stFileUploader"]{
    background:var(--panel) !important;
    border:1px dashed var(--line) !important;
    border-radius:14px !important;
    padding:.55rem .7rem !important;
    box-shadow:none !important;
}

/* dataframe */
[data-testid="stDataFrame"]{
    border:1px solid var(--line) !important;
    border-radius:12px !important;
    overflow:hidden !important;
    box-shadow:none !important;
}

/* horizontal layout */
[data-testid="stHorizontalBlock"]{
    align-items:flex-start !important;
    gap:1.1rem !important;
}

/* 工具列那排：縮小間距＋不允許換行 */
[data-testid="stHorizontalBlock"]:has(> [data-testid="stColumn"] > [data-testid="stPopover"]) {
    gap: 0.35rem !important;
    flex-wrap: nowrap !important;
    align-items: center !important;
}
[data-testid="stHorizontalBlock"]:has(> [data-testid="stColumn"] > [data-testid="stPopover"]) button {
    font-size: 13px !important;
    padding-left: 8px !important;
    padding-right: 8px !important;
    white-space: nowrap !important;
}

hr{
    border:none !important;
    border-top:1px solid rgba(201,179,148,.55) !important;
    margin:1.2rem 0 !important;
}
</style>
""", unsafe_allow_html=True)


st.markdown("""
<style>
.toolbar-row{margin-top:.35rem;}
.small-note{color:#8a7e71;font-size:12px;}

/* popover 不要把工具列撐爆 */
[data-testid="stPopover"] > div{
    min-width: unset !important;
    width: max-content !important;
    max-width: min(92vw, 760px) !important;
}

/* 工具列的 popover 按鈕填滿欄寬 */
[data-testid="stColumn"] > [data-testid="stPopover"] {
    width: 100% !important;
}
[data-testid="stColumn"] > [data-testid="stPopover"] > button {
    width: 100% !important;
    white-space: nowrap !important;
}

/* 無 Ace 時的列號欄 */
.line-no-wrap textarea{
    text-align:right !important;
    color:#8a7e71 !important;
    background:#f5f0e8 !important;
    border-right:none !important;
    padding-right:10px !important;
    font-family:Consolas, monospace !important;
    font-size:14px !important;
    line-height:1.72 !important;
}
.line-no-wrap [data-baseweb="textarea"]{
    border-top-right-radius:0 !important;
    border-bottom-right-radius:0 !important;
}
.editor-main-wrap [data-baseweb="textarea"]{
    border-top-left-radius:0 !important;
    border-bottom-left-radius:0 !important;
}
</style>
<script>
(function(){
    function disableWrapIndent(){
        if(!window.ace) return false;
        var found = false;
        document.querySelectorAll('.ace_editor').forEach(function(el){
            try{ window.ace.edit(el).session.setOption('indentedSoftWrap',false); found=true; }catch(e){}
        });
        return found;
    }
    function tryUntilReady(n){
        if(n<=0) return;
        if(!disableWrapIndent()) setTimeout(function(){tryUntilReady(n-1);},300);
        else new MutationObserver(disableWrapIndent).observe(document.body,{childList:true,subtree:true});
    }
    tryUntilReady(20);
})();
</script>
""", unsafe_allow_html=True)

for key, default in [
    ("title", ""), ("time", ""), ("chair", ""), ("location", ""),
    ("recorder", ""), ("attendees", ""), ("highest", "副總經理"),
    ("content", ""), ("owner_text", ""), ("attachments", []), ("search_kw", ""),
    ("show_word_like_preview", False),
    ("show_attachment_panel", False),
    ("generated_doc_bytes", None), ("generated_doc_name", ""),
    ("preview_open", False),
]:
    if key not in st.session_state:
        st.session_state[key] = default

if "_all_records" not in st.session_state:
    st.session_state["_all_records"] = []
if "editor_rev" not in st.session_state:
    st.session_state["editor_rev"] = 0
if "_last_normalized_content" not in st.session_state:
    st.session_state["_last_normalized_content"] = st.session_state.get("content", "")

sync_owner_rows()
refresh_row_map_cache()


def _record_store_dir():
    folder = os.path.join(os.path.expanduser("~"), "Desktop", "會議記錄助手")
    os.makedirs(folder, exist_ok=True)
    return folder


def _sanitize_record_filename(text: str) -> str:
    text = (text or "").strip()
    text = re.sub(r'[\\/:*?"<>|\r\n]+', "_", text)
    text = re.sub(r"\s+", " ", text).strip(" ._")
    return text or "未命名"


def _make_record_label(rec):
    title = rec.get("title", "（無標題）") or "（無標題）"
    time_ = rec.get("time", "") or ""
    saved_at = rec.get("saved_at", "") or ""
    return f"{title}　{time_}　[{saved_at}]".strip() if saved_at else f"{title}　{time_}".strip()


def _load_all_records():
    folder = _record_store_dir()
    records = []
    for name in sorted(os.listdir(folder), reverse=True):
        if not name.lower().endswith(".json"):
            continue
        full = os.path.join(folder, name)
        try:
            with open(full, "r", encoding="utf-8") as f:
                rec = json.load(f)
            rec["_file_name"] = name
            rec["_file_path"] = full
            records.append(rec)
        except Exception:
            continue
    return records


def _delete_record_file(rec):
    full = rec.get("_file_path")
    if full and os.path.exists(full):
        os.remove(full)


def _save_all_records(records):
    # 本機私有記錄模式：不做整批覆寫，保留函式名稱相容舊呼叫。
    return None


def _build_current_record():
    return {
        "title": st.session_state.get("title", ""),
        "time": st.session_state.get("time", ""),
        "chair": st.session_state.get("chair", ""),
        "location": st.session_state.get("location", ""),
        "recorder": st.session_state.get("recorder", ""),
        "attendees": st.session_state.get("attendees", ""),
        "highest": st.session_state.get("highest", "副總經理"),
        "content": st.session_state.get("content", ""),
        "owner_text": st.session_state.get("owner_text", ""),
        "attachments": st.session_state.get("attachments", []),
    }


def _save_current_to_records():
    folder = _record_store_dir()
    rec = _build_current_record()
    rec["saved_at"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    safe_title = _sanitize_record_filename(rec.get("title", ""))
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    file_name = f"{ts}_{safe_title}.json"
    full = os.path.join(folder, file_name)
    with open(full, "w", encoding="utf-8") as f:
        json.dump(rec, f, ensure_ascii=False, indent=2)
    return full



left_col, center_col, right_col = st.columns([0.62, 1.18, 0.48], gap="medium")

with left_col:
    st.subheader("基本資料")
    st.session_state["title"] = st.text_input("會議名稱", st.session_state["title"])
    st.session_state["time"] = st.text_input("開會時間", st.session_state["time"])
    st.session_state["chair"] = st.text_input("主　　席", st.session_state["chair"])
    st.session_state["location"] = st.text_input("開會地點", st.session_state["location"])
    st.session_state["recorder"] = st.text_input("會議紀錄", st.session_state["recorder"])
    st.session_state["attendees"] = st.text_area("出列席人員（換行輸入）", st.session_state["attendees"], height=120)
    st.session_state["highest"] = st.selectbox(
        "最高簽核主管",
        ["副總經理", "總經理"],
        index=0 if st.session_state["highest"] == "副總經理" else 1
    )

    st.markdown("---")
    st.caption("記錄管理")
    if st.button("💾 儲存目前記錄", use_container_width=True):
        _saved_path = _save_current_to_records()
        st.success(f"已儲存到：{_saved_path}")

    _records = _load_all_records()
    if _records:
        _labels = [_make_record_label(r) for r in _records]
        _sel = st.selectbox("選擇記錄", _labels, key="record_selector", label_visibility="collapsed")
        load_col, del_col = st.columns(2)
        if load_col.button("📥 載入", use_container_width=True):
            for rec in _records:
                if _make_record_label(rec) == _sel:
                    st.session_state["title"] = rec.get("title", "")
                    st.session_state["time"] = rec.get("time", "")
                    sync_time_picker_from_text(st.session_state["time"])
                    st.session_state["chair"] = rec.get("chair", "")
                    st.session_state["location"] = rec.get("location", "")
                    st.session_state["recorder"] = rec.get("recorder", "")
                    st.session_state["attendees"] = rec.get("attendees", "")
                    st.session_state["highest"] = rec.get("highest", "副總經理")
                    st.session_state["content"] = rec.get("content", "")
                    st.session_state["owner_text"] = rec.get("owner_text", "")
                    st.session_state["attachments"] = rec.get("attachments", [])
                    st.session_state["search_kw"] = ""
                    st.session_state["generated_doc_bytes"] = None
                    st.session_state["generated_doc_name"] = ""
                    st.session_state["editor_rev"] += 1
                    st.rerun()
        if del_col.button("🗑️ 刪除", use_container_width=True):
            for rec in _records:
                if _make_record_label(rec) == _sel:
                    _delete_record_file(rec)
                    break
            st.rerun()
    else:
        st.caption("尚無儲存的記錄")

    st.caption(f"首次儲存時，會自動在桌面建立：{os.path.join(os.path.expanduser('~'), 'Desktop', '會議記錄助手')}")

with center_col:
    st.subheader("會議記錄內容")

    editor_value = build_editor_display(
        st.session_state.get("content", ""),
        st.session_state.get("attachments", [])
    )

    editor_height = 300
    editor_key = f"content_main_area_{st.session_state['editor_rev']}"

    if HAS_ACE:
        content_tmp = st_ace(
            value=editor_value,
            language="text",
            theme="textmate",
            key=editor_key,
            height=editor_height,
            font_size=18,
            tab_size=2,
            wrap=True,
            show_gutter=True,
            show_print_margin=False,
            auto_update=True,
            placeholder="例：逐項記錄會議討論內容",
        )
    else:
        editor_lines = editor_value.replace("\r\n", "\n").replace("\r", "\n").split("\n")
        line_count = max(1, len(editor_lines))
        line_numbers = "\n".join(str(i) for i in range(1, line_count + 1))
        ln_col, ed_col = st.columns([0.10, 0.90], gap="none")
        with ln_col:
            st.markdown('<div class="line-no-wrap">', unsafe_allow_html=True)
            st.text_area(
                "列號",
                value=line_numbers,
                height=editor_height,
                key=f"{editor_key}_line_numbers",
                disabled=True,
                label_visibility="collapsed",
            )
            st.markdown('</div>', unsafe_allow_html=True)
        with ed_col:
            st.markdown('<div class="editor-main-wrap">', unsafe_allow_html=True)
            content_tmp = st.text_area(
                "會議記錄內容",
                value=editor_value,
                height=editor_height,
                key=editor_key,
                label_visibility="collapsed",
                placeholder="例：逐項記錄會議討論內容",
            )
            st.markdown('</div>', unsafe_allow_html=True)

    content_tmp = content_tmp if content_tmp is not None else editor_value
    pure, tampered = strip_placeholders(
        content_tmp,
        get_expected_placeholders(st.session_state.get("attachments", []))
    )
    if tampered:
        st.session_state["_att_tampered"] = True
        st.session_state["editor_rev"] += 1
        st.rerun()

    if pure != st.session_state.get("content", ""):
        st.session_state["content"] = pure
        sync_owner_rows()
        refresh_row_map_cache()

    if st.session_state.get("_att_tampered"):
        st.warning("⚠️ 附件佔位符被修改，已自動還原。附件列不可手動編輯，請至附件清單刪除附件。")
        st.session_state["_att_tampered"] = False

    cache = get_row_map_cache()
    normalized_content = normalize_content_for_editor(st.session_state.get("content", ""))
    expanded = expand_content_lines(normalized_content)
    owner_target_rows = len(expanded)
    owner_value = normalize_owner_text(st.session_state.get("owner_text", ""), owner_target_rows)
    if owner_value != st.session_state.get("owner_text", ""):
        st.session_state["owner_text"] = owner_value

    stats = calc_page_stats(st.session_state["content"], st.session_state["attachments"])
    page_label = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"][min(stats["page_num"] - 1, 9)]

    tool1, tool2, tool3, tool4, tool5 = st.columns([1.05, 1.25, 1.25, 1.0, 1.45], gap="small")

    with tool1:
        do_format = st.button("📝 排版", use_container_width=True)

    with tool2:
        with st.popover("🔍 搜尋", use_container_width=True):
            search_kw = st.text_input("搜尋關鍵字", key="search_kw", placeholder="例：逆境")
            if search_kw.strip():
                hits = find_matching_lines(search_kw.strip())
                if hits:
                    row_text = "、".join([str(row_no) for row_no, _ in hits])
                    st.success(f"共找到 {len(hits)} 列：第 {row_text} 列")
                else:
                    st.info("找不到符合的內容")

    with tool3:
        with st.popover("📎 附件", use_container_width=True):
            att_tab1, att_tab2, att_tab3 = st.tabs(["＋ 新增圖片", "＋ 手動輸入表格", "＋ 匯入 xlsx"])

            existing_img_labels = set()
            existing_table_labels = set()
            for a in st.session_state["attachments"]:
                if a["type"] == "image":
                    existing_img_labels.add(a.get("label", ""))
                elif a["type"] == "image_pair":
                    existing_img_labels.update(a.get("member_labels", []))
                elif a["type"] == "table":
                    existing_table_labels.add(a.get("label", ""))

            with att_tab1:
                img_files = st.file_uploader("選擇圖片", type=["png", "jpg", "jpeg", "bmp", "gif", "webp"], key="img_upload", accept_multiple_files=True)
                r1c1, r1c2 = st.columns(2)
                img_after = r1c1.number_input("插入在第幾行後（0=最前）", 0, 999, max(0, cache["visual_count"]), key="img_after")
                img_cost = r1c2.number_input("每組估幾列", 1, 30, 6, key="img_cost")
                r2c1, r2c2 = st.columns(2)
                img_cap = r2c1.text_input("附件總標題（可留空）", key="img_cap")
                img_layout = r2c2.selectbox("排版方式", ["單張", "左右各一張"], key="img_layout")
                img_cappos = st.radio("總標題位置", ["上方", "下方"], horizontal=True, key="img_cappos")

                pair_cap_left = ""
                pair_cap_right = ""
                if img_layout == "左右各一張":
                    pc1, pc2 = st.columns(2)
                    pair_cap_left = pc1.text_input("左圖標題", key="pair_cap_left")
                    pair_cap_right = pc2.text_input("右圖標題", key="pair_cap_right")

                if st.button("加入附件清單", key="btn_add_img_form", use_container_width=False):
                    if not img_files:
                        st.warning("請先選擇圖片。")
                    else:
                        duplicated = [f.name for f in img_files if f.name in existing_img_labels]
                        added = 0
                        skipped = duplicated[:]
                        usable_files = [f for f in img_files if f.name not in existing_img_labels]
                        try:
                            if img_layout == "單張":
                                for img_file in usable_files:
                                    tmp_path = save_uploaded_file_to_temp(img_file)
                                    st.session_state["attachments"].append({"type": "image", "path": tmp_path, "after_line": int(img_after), "row_cost": int(img_cost), "caption": img_cap.strip(), "caption_pos": "above" if img_cappos == "上方" else "below", "label": img_file.name})
                                    added += 1
                            else:
                                if len(usable_files) < 2:
                                    st.warning("左右各一張模式至少要選 2 張圖片。")
                                else:
                                    if len(usable_files) % 2 == 1:
                                        skipped.append(usable_files[-1].name)
                                        usable_files = usable_files[:-1]
                                    for i in range(0, len(usable_files), 2):
                                        left_file = usable_files[i]
                                        right_file = usable_files[i + 1]
                                        left_path = save_uploaded_file_to_temp(left_file)
                                        right_path = save_uploaded_file_to_temp(right_file)
                                        st.session_state["attachments"].append({"type": "image_pair", "left_path": left_path, "right_path": right_path, "after_line": int(img_after), "row_cost": int(img_cost), "caption": img_cap.strip(), "left_caption": pair_cap_left.strip(), "right_caption": pair_cap_right.strip(), "caption_pos": "above" if img_cappos == "上方" else "below", "label": f"{left_file.name}｜{right_file.name}", "member_labels": [left_file.name, right_file.name]})
                                        added += 1
                        except Exception as e:
                            st.error(f"加入圖片附件失敗：{e}")
                            added = 0
                        if added:
                            clear_attachment_upload_widgets()
                            st.session_state["editor_rev"] += 1
                            st.rerun()
                        elif skipped:
                            st.info(f"已略過重複或未成對檔案：{', '.join(skipped)}")

            with att_tab2:
                if not HAS_PANDAS:
                    st.warning("請先安裝 pandas：pip install pandas")
                else:
                    c1, c2, c3, c4 = st.columns(4)
                    tbl_after = c1.number_input("插入在第幾行後（0=最前）", 0, 999, max(0, stats["total_rows"]), key="tbl_after")
                    tbl_cols = int(c2.number_input("欄數", 1, 10, 3, key="tbl_cols"))
                    tbl_rows = int(c3.number_input("列數（含標題列）", 2, 30, 4, key="tbl_rows"))
                    tbl_row_cost = c4.number_input("每組估幾列", 1, 40, max(4, tbl_rows + 2), key="tbl_row_cost")
                    tbl_cap = c1.text_input("標題文字（可留空）", key="tbl_cap")
                    tbl_cappos = st.radio("標題位置", ["上方", "下方"], horizontal=True, key="tbl_cappos")
                    st.markdown("**填寫表格內容（第一列為標題列）：**")

                    df_state_key = f"tbl_df_{tbl_rows}_{tbl_cols}"
                    if df_state_key not in st.session_state:
                        st.session_state[df_state_key] = pd.DataFrame([[""] * tbl_cols for _ in range(tbl_rows)], columns=[f"欄{i + 1}" for i in range(tbl_cols)])
                    edited_df = st.data_editor(st.session_state[df_state_key], use_container_width=True, key=f"tbl_editor_{tbl_rows}_{tbl_cols}", hide_index=True)
                    st.session_state[df_state_key] = edited_df.copy()

                    if st.button("加入附件清單", key="btn_add_tbl"):
                        try:
                            rows_data = edited_df.fillna("").astype(str).values.tolist()
                            st.session_state["attachments"].append({"type": "table", "data": rows_data, "after_line": int(tbl_after), "row_cost": int(tbl_row_cost), "caption": tbl_cap.strip(), "caption_pos": "above" if tbl_cappos == "上方" else "below", "label": f"手動表格 {tbl_rows}×{tbl_cols}"})
                            clear_attachment_upload_widgets()
                            st.session_state["editor_rev"] += 1
                            st.rerun()
                        except Exception as e:
                            st.error(f"加入手動表格失敗：{e}")

            with att_tab3:
                if not HAS_OPENPYXL:
                    st.warning("請先安裝 openpyxl：pip install openpyxl")
                else:
                    xlsx_files = st.file_uploader("選擇 Excel 檔案", type=["xlsx", "xls"], key="xlsx_upload", accept_multiple_files=True)
                    if xlsx_files:
                        try:
                            preview_idx = st.selectbox("先選擇要處理的檔案", options=range(len(xlsx_files)), format_func=lambda i: xlsx_files[i].name, key="xlsx_preview_idx")
                            xlsx_file = xlsx_files[preview_idx]
                            xlsx_bytes = uploaded_file_bytes(xlsx_file)
                            wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
                            c1, c2, c3 = st.columns(3)
                            xlsx_after = c1.number_input("插入在第幾行後（0=最前）", 0, 999, max(0, cache["visual_count"]), key="xlsx_after")
                            xlsx_sheet = c2.selectbox("工作表", wb.sheetnames, key="xlsx_sheet")
                            xlsx_row_cost = c3.number_input("每組估幾列", 1, 80, 8, key="xlsx_row_cost")
                            xlsx_cap = st.text_input("標題文字（可留空）", key="xlsx_cap")
                            xlsx_cappos = st.radio("標題位置", ["上方", "下方"], horizontal=True, key="xlsx_cappos")
                            candidate_label = f"{xlsx_file.name} [{xlsx_sheet}]"
                            if candidate_label in existing_table_labels:
                                st.warning(f"「{candidate_label}」已在附件清單中，請先刪除舊的再重新加入。")
                            elif st.button("加入附件清單", key="btn_add_xlsx"):
                                ws = wb[xlsx_sheet]
                                data = []
                                for row in ws.iter_rows(values_only=True):
                                    if any(cell is not None for cell in row):
                                        data.append([str(cell) if cell is not None else "" for cell in row])
                                if not data:
                                    st.error("工作表內沒有資料")
                                else:
                                    st.session_state["attachments"].append({"type": "table", "data": data, "after_line": int(xlsx_after), "row_cost": int(xlsx_row_cost), "caption": xlsx_cap.strip(), "caption_pos": "above" if xlsx_cappos == "上方" else "below", "label": candidate_label})
                                    clear_attachment_upload_widgets()
                                    st.session_state["editor_rev"] += 1
                                    st.rerun()
                        except Exception as e:
                            st.error(f"無法開啟 Excel：{e}")

            if st.session_state["attachments"]:
                st.markdown("---")
                st.markdown("**📋 附件清單**")
                for i, att in enumerate(st.session_state["attachments"]):
                    icon = "🖼🖼" if att["type"] == "image_pair" else ("🖼" if att["type"] == "image" else "📊")
                    c1, c2, c3, c4 = st.columns([3, 2, 2, 1])
                    c1.markdown(f"{icon} `{att['label']}`")
                    c2.markdown(f"第 **{att['after_line']}** 行後・佔 **{att['row_cost']}** 列")
                    if att.get("type") == "image_pair" and (att.get("left_caption") or att.get("right_caption")):
                        c3.markdown(f"左右標題：{att.get('left_caption','')}｜{att.get('right_caption','')}")
                    elif att.get("caption"):
                        c3.markdown(f"標題：{att['caption']}")
                    if c4.button("✕", key=f"del_att_{i}"):
                        if att["type"] == "image" and att.get("path"):
                            try: os.unlink(att["path"])
                            except Exception: pass
                        elif att["type"] == "image_pair":
                            for p in [att.get("left_path"), att.get("right_path")]:
                                if p:
                                    try: os.unlink(p)
                                    except Exception: pass
                        st.session_state["attachments"].pop(i)
                        st.session_state["editor_rev"] += 1
                        clear_attachment_upload_widgets()
                        st.rerun()

    with tool4:
        open_preview = st.button("👁️ 預覽", use_container_width=True)

    with tool5:
        generate_now = st.button("📄 產生 Word", type="primary", use_container_width=True)

    st.caption(f"編輯器左側列號可直接對照搜尋結果｜目前 {stats['total_rows']} 列・第{page_label}頁 {stats['rows_in_page']} 列")

    if do_format:
        formatted = format_content_for_word_button(st.session_state.get("content", ""))
        if formatted != st.session_state.get("content", ""):
            st.session_state["content"] = formatted
            sync_owner_rows()
            refresh_row_map_cache()
            st.session_state["editor_rev"] += 1
            st.rerun()

    if open_preview:
        st.session_state["preview_open"] = True

    if generate_now:
        data = {"title": st.session_state["title"], "time": st.session_state["time"], "chair": st.session_state["chair"], "location": st.session_state["location"], "recorder": st.session_state["recorder"], "attendees": st.session_state["attendees"], "highest": st.session_state["highest"]}
        _save_current_to_records()
        try:
            buf = generate_doc_bytes(data, st.session_state["content"], st.session_state["owner_text"], st.session_state["attachments"])
            st.session_state["generated_doc_bytes"] = buf.getvalue()
            st.session_state["generated_doc_name"] = f"會議紀錄_{data['title'] or '未命名'}.docx"
            st.success("✅ 產生成功，請點下方按鈕下載。")
        except FileNotFoundError as e:
            st.error(str(e))
        except Exception as e:
            import traceback
            st.error(f"❌ 產生失敗：{e}")
            with st.expander("詳細錯誤"):
                st.code(traceback.format_exc())

    if st.session_state.get("generated_doc_bytes"):
        st.download_button(label="⬇️ 下載 Word 檔", data=st.session_state["generated_doc_bytes"], file_name=st.session_state.get("generated_doc_name", "會議紀錄.docx"), mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

with right_col:
    st.subheader("決議執行者")
    owner_tmp = st.text_area("決議執行者", value=owner_value, height=300, key=f"owner_area_{owner_target_rows}", label_visibility="collapsed", placeholder="逐列輸入執行者")
    normalized_owner = normalize_owner_text(owner_tmp, owner_target_rows)
    if normalized_owner != st.session_state.get("owner_text", ""):
        st.session_state["owner_text"] = normalized_owner

if st.session_state.get("preview_open", False):
    @st.dialog("📄 文件預覽", width="large")
    def _preview_dialog():
        if st.button("✕ 關閉預覽", use_container_width=True):
            st.session_state["preview_open"] = False
            st.rerun()
        render_word_like_preview(st.session_state.get("content", ""), st.session_state.get("owner_text", ""), st.session_state.get("attachments", []))
    _preview_dialog()